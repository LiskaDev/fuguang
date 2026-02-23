"""
FileGeneratorSkills — 📁 文件生成类技能
生成多种格式文件（文本/CSV/Excel/Word/PDF）并推送到 Web UI 供用户下载。

所有工具统一注册文件到 _pending_file_cards 队列，
由 web_bridge.py 拦截后注册到 _files 并推送 WebSocket 下载卡片。
"""
import os
import csv
import json
import logging
from pathlib import Path

logger = logging.getLogger("fuguang.skills")


# ========================
# 📋 工具 Schema
# ========================
_FILE_GENERATOR_TOOLS_SCHEMA = [
    {"type": "function", "function": {
        "name": "generate_text_file",
        "description": (
            "生成文本类文件供用户在网页端下载。"
            "支持任何文本格式：.txt, .md, .py, .js, .html, .css, .json, .xml, .yaml, .sql 等。\n"
            "【格式选择规则】用户要求哪种格式就用哪种工具，绝对不要替换格式。此工具仅用于纯文本类文件。\n"
            "要PDF→generate_pdf，要Excel→generate_xlsx，要Word→generate_docx，要CSV→generate_csv。\n"
            "⚠️ 此工具用于将内容打包为文件供下载，与 create_file_directly（写入本地磁盘）不同。"
        ),
        "parameters": {"type": "object", "properties": {
            "filename": {"type": "string", "description": "文件名（含扩展名），如 'report.md'、'app.py'"},
            "content": {"type": "string", "description": "文件内容"}
        }, "required": ["filename", "content"]}
    }},
    {"type": "function", "function": {
        "name": "generate_csv",
        "description": (
            "生成CSV文件供用户下载。适用于表格数据导出。\n"
            "【格式选择规则】用户要求哪种格式就用哪种工具，绝对不要替换。\n"
            "用户说CSV用此工具，说Excel用generate_xlsx，不要混用。\n"
            "示例：generate_csv(filename='sales.csv', headers=['日期','金额'], "
            "rows=[['2024-01','1000'], ['2024-02','2000']])"
        ),
        "parameters": {"type": "object", "properties": {
            "filename": {"type": "string", "description": "文件名，如 'data.csv'"},
            "headers": {"type": "array", "items": {"type": "string"}, "description": "表头列名"},
            "rows": {"type": "array", "items": {"type": "array", "items": {"type": "string"}},
                     "description": "数据行（二维数组），每行元素数需与表头一致"}
        }, "required": ["filename", "headers", "rows"]}
    }},
    {"type": "function", "function": {
        "name": "generate_xlsx",
        "description": (
            "生成Excel(.xlsx)文件供用户下载。支持多个Sheet页。\n"
            "【格式选择规则】用户要求哪种格式就用哪种工具。\n"
            "用户说Excel/表格用此工具，说CSV用generate_csv，不要混用。\n"
            "⚠️ 需要 openpyxl 库（无则自动降级为CSV）。"
        ),
        "parameters": {"type": "object", "properties": {
            "filename": {"type": "string", "description": "文件名，如 'report.xlsx'"},
            "sheets": {"type": "array", "items": {"type": "object", "properties": {
                "name": {"type": "string", "description": "Sheet名称"},
                "headers": {"type": "array", "items": {"type": "string"}, "description": "表头"},
                "rows": {"type": "array", "items": {"type": "array"}, "description": "数据行"}
            }, "required": ["name", "headers", "rows"]}, "description": "Sheet页列表"}
        }, "required": ["filename", "sheets"]}
    }},
    {"type": "function", "function": {
        "name": "generate_docx",
        "description": (
            "生成Word(.docx)文件供用户下载。支持标题、正文段落和表格。\n"
            "【格式选择规则】用户要求哪种格式就用哪种工具。\n"
            "仅当用户明确说\"Word\"、\"docx\"时使用。说PDF用generate_pdf，说Excel用generate_xlsx。\n"
            "⚠️ 需要 python-docx 库。"
        ),
        "parameters": {"type": "object", "properties": {
            "filename": {"type": "string", "description": "文件名，如 'report.docx'"},
            "title": {"type": "string", "description": "文档标题"},
            "blocks": {"type": "array", "items": {"type": "object", "properties": {
                "type": {"type": "string", "enum": ["heading", "paragraph", "table"],
                         "description": "内容块类型"},
                "level": {"type": "integer",
                          "description": "标题级别(1-4)，仅 heading 类型需要"},
                "text": {"type": "string",
                         "description": "文本内容，heading 和 paragraph 类型需要"},
                "headers": {"type": "array", "items": {"type": "string"},
                            "description": "表头，仅 table 类型需要"},
                "rows": {"type": "array", "items": {"type": "array", "items": {"type": "string"}},
                         "description": "数据行，仅 table 类型需要"}
            }, "required": ["type"]}, "description": "文档内容块列表"}
        }, "required": ["filename", "title", "blocks"]}
    }},
    {"type": "function", "function": {
        "name": "generate_pdf",
        "description": (
            "生成PDF文件供用户下载。支持中文，适合正式文档和报告。\n"
            "【格式选择规则】用户要求哪种格式就用哪种工具。\n"
            "用户说PDF/报告/文档（未指定格式）用此工具。说Word用generate_docx，说Excel用generate_xlsx。\n"
            "⚠️ 需要 fpdf2 库。中文使用系统微软雅黑字体。"
        ),
        "parameters": {"type": "object", "properties": {
            "filename": {"type": "string", "description": "文件名，如 'report.pdf'"},
            "title": {"type": "string", "description": "文档标题"},
            "content": {"type": "string",
                        "description": "正文内容，用换行符分段"}
        }, "required": ["filename", "title", "content"]}
    }},
]


class FileGeneratorSkills:
    """文件生成类技能 Mixin — 生成文件并推送下载卡片"""
    _FILE_GENERATOR_TOOLS = _FILE_GENERATOR_TOOLS_SCHEMA

    # ========================
    # 🔧 内部工具
    # ========================

    def _get_temp_files_dir(self) -> Path:
        """获取临时文件目录（自动创建）"""
        temp_dir = self.config.PROJECT_ROOT / "temp_files"
        temp_dir.mkdir(parents=True, exist_ok=True)
        return temp_dir

    def _register_file_card(self, filepath: str, filename: str):
        """注册文件卡片到待推送队列（由 web_bridge 拦截并推送）"""
        if not hasattr(self, '_pending_file_cards'):
            self._pending_file_cards = []
        try:
            size = os.path.getsize(filepath)
        except Exception:
            size = 0
        self._pending_file_cards.append({
            "filepath": filepath,
            "filename": filename,
            "size": size
        })
        logger.info(f"📁 [文件生成] 已注册下载卡片: {filename} ({size} bytes)")

    def _safe_filename(self, filename: str) -> str:
        """防止路径穿越，只取文件名部分"""
        return Path(filename).name

    # ========================
    # 📄 文本文件
    # ========================

    def generate_text_file(self, filename: str, content: str) -> str:
        """
        生成文本类文件供下载。支持 .txt/.md/.py/.js/.html/.css/.json 等任何文本格式。

        Args:
            filename: 文件名（含扩展名）
            content: 文件内容
        """
        try:
            safe_name = self._safe_filename(filename)
            filepath = str(self._get_temp_files_dir() / safe_name)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)
            self._register_file_card(filepath, safe_name)
            return f"✅ 文件已生成: {safe_name}，已推送下载链接"
        except Exception as e:
            logger.error(f"📁 [文件生成] 文本文件失败: {e}")
            return f"❌ 生成文件失败: {str(e)}"

    # ========================
    # 📊 CSV 文件
    # ========================

    def generate_csv(self, filename: str, headers: list, rows: list) -> str:
        """
        生成CSV文件供下载。

        Args:
            filename: 文件名
            headers: 表头列名列表
            rows: 数据行（二维数组）
        """
        try:
            safe_name = self._safe_filename(filename)
            if not safe_name.endswith('.csv'):
                safe_name += '.csv'
            filepath = str(self._get_temp_files_dir() / safe_name)
            with open(filepath, "w", encoding="utf-8-sig", newline="") as f:
                writer = csv.writer(f)
                writer.writerow(headers)
                writer.writerows(rows)
            self._register_file_card(filepath, safe_name)
            return f"✅ CSV 文件已生成: {safe_name}（{len(rows)}行），已推送下载链接"
        except Exception as e:
            logger.error(f"📁 [文件生成] CSV 失败: {e}")
            return f"❌ 生成CSV失败: {str(e)}"

    # ========================
    # 📗 Excel 文件
    # ========================

    def generate_xlsx(self, filename: str, sheets: list) -> str:
        """
        生成Excel(.xlsx)文件供下载，支持多Sheet。

        Args:
            filename: 文件名
            sheets: Sheet页列表，每个含 name/headers/rows
        """
        try:
            import openpyxl
        except ImportError:
            logger.warning("📁 [文件生成] openpyxl 未安装，降级为 CSV")
            # 降级：取第一个 Sheet 生成 CSV
            if sheets:
                s = sheets[0]
                csv_name = filename.replace('.xlsx', '.csv')
                return self.generate_csv(csv_name, s.get("headers", []), s.get("rows", []))
            return "❌ Excel生成功能暂不可用（缺少openpyxl），已尝试降级但无数据"

        try:
            safe_name = self._safe_filename(filename)
            if not safe_name.endswith('.xlsx'):
                safe_name += '.xlsx'
            filepath = str(self._get_temp_files_dir() / safe_name)

            wb = openpyxl.Workbook()
            # 删除默认 Sheet
            wb.remove(wb.active)

            for sheet_data in sheets:
                ws = wb.create_sheet(title=sheet_data.get("name", "Sheet"))
                headers = sheet_data.get("headers", [])
                rows = sheet_data.get("rows", [])

                # 写入表头（加粗）
                from openpyxl.styles import Font
                for col_idx, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col_idx, value=header)
                    cell.font = Font(bold=True)

                # 写入数据
                for row_idx, row in enumerate(rows, 2):
                    for col_idx, value in enumerate(row, 1):
                        ws.cell(row=row_idx, column=col_idx, value=value)

                # 自动列宽
                for col in ws.columns:
                    max_len = max(len(str(cell.value or "")) for cell in col)
                    ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

            wb.save(filepath)
            total_rows = sum(len(s.get("rows", [])) for s in sheets)
            self._register_file_card(filepath, safe_name)
            return f"✅ Excel 文件已生成: {safe_name}（{len(sheets)}个Sheet，共{total_rows}行），已推送下载链接"
        except Exception as e:
            logger.error(f"📁 [文件生成] Excel 失败: {e}")
            return f"❌ 生成Excel失败: {str(e)}"

    # ========================
    # 📝 Word 文件
    # ========================

    def generate_docx(self, filename: str, title: str, blocks: list) -> str:
        """
        生成Word(.docx)文件供下载，支持标题/正文/表格。

        Args:
            filename: 文件名
            title: 文档标题
            blocks: 内容块列表，每块含 type(heading/paragraph/table) + 对应字段
        """
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return "❌ Word生成功能暂不可用（缺少python-docx库），请告知用户此功能暂时无法使用"

        try:
            safe_name = self._safe_filename(filename)
            if not safe_name.endswith('.docx'):
                safe_name += '.docx'
            filepath = str(self._get_temp_files_dir() / safe_name)

            doc = Document()
            doc.add_heading(title, level=0)

            for block in blocks:
                block_type = block.get("type", "paragraph")

                if block_type == "heading":
                    level = block.get("level", 1)
                    doc.add_heading(block.get("text", ""), level=min(level, 4))

                elif block_type == "paragraph":
                    doc.add_paragraph(block.get("text", ""))

                elif block_type == "table":
                    headers = block.get("headers", [])
                    rows = block.get("rows", [])
                    if headers:
                        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                        table.style = 'Table Grid'
                        # 表头
                        for i, h in enumerate(headers):
                            table.rows[0].cells[i].text = str(h)
                        # 数据行
                        for row_idx, row in enumerate(rows):
                            for col_idx, val in enumerate(row):
                                if col_idx < len(headers):
                                    table.rows[row_idx + 1].cells[col_idx].text = str(val)

            doc.save(filepath)
            self._register_file_card(filepath, safe_name)
            return f"✅ Word 文件已生成: {safe_name}，已推送下载链接"
        except Exception as e:
            logger.error(f"📁 [文件生成] Word 失败: {e}")
            return f"❌ 生成Word失败: {str(e)}"

    # ========================
    # 📕 PDF 文件
    # ========================

    def generate_pdf(self, filename: str, title: str, content: str) -> str:
        """
        生成PDF文件供下载，支持中文（使用系统微软雅黑字体）。

        Args:
            filename: 文件名
            title: 文档标题
            content: 正文内容，换行符分段
        """
        try:
            from fpdf import FPDF
        except ImportError:
            return "❌ PDF生成功能暂不可用（缺少fpdf2库），请告知用户此功能暂时无法使用"

        try:
            safe_name = self._safe_filename(filename)
            if not safe_name.endswith('.pdf'):
                safe_name += '.pdf'
            filepath = str(self._get_temp_files_dir() / safe_name)

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # 加载中文字体
            font_loaded = False
            font_paths = [
                "C:/Windows/Fonts/msyh.ttc",   # 微软雅黑
                "C:/Windows/Fonts/simsun.ttc",  # 宋体
                "C:/Windows/Fonts/simhei.ttf",  # 黑体
            ]
            for fp in font_paths:
                if os.path.exists(fp):
                    try:
                        pdf.add_font("CJK", "", fp, uni=True)
                        pdf.set_font("CJK", size=12)
                        font_loaded = True
                        break
                    except Exception as e:
                        logger.debug(f"📁 [PDF] 字体加载失败 {fp}: {e}")
                        continue

            if not font_loaded:
                # 降级：使用内置字体（不支持中文）
                pdf.set_font("Helvetica", size=12)
                logger.warning("📁 [PDF] 未找到中文字体，降级为 Helvetica")

            # 标题
            if font_loaded:
                pdf.set_font("CJK", size=18)
            else:
                pdf.set_font("Helvetica", "B", 18)
            pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align="C")
            pdf.ln(8)

            # 正文
            if font_loaded:
                pdf.set_font("CJK", size=12)
            else:
                pdf.set_font("Helvetica", size=12)

            for paragraph in content.split("\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    pdf.multi_cell(0, 7, paragraph)
                    pdf.ln(3)

            pdf.output(filepath)
            self._register_file_card(filepath, safe_name)
            return f"✅ PDF 文件已生成: {safe_name}，已推送下载链接"
        except Exception as e:
            logger.error(f"📁 [文件生成] PDF 失败: {e}")
            return f"❌ 生成PDF失败: {str(e)}"
